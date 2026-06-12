from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\13784\Desktop\git\NavTools")
OUTPUT = ROOT / "docs" / "NewAPI中转站配置OpenClaw小龙虾用户操作手册.docx"

BLUE = "1F4D78"
TITLE_BLUE = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "B8C7D9"
RED = "C00000"
MUTED = "555555"


def add_border(element, color=GRID, size="8"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    element._tbl.tblPr.append(borders)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def write_runs(paragraph, parts, default_size=11):
    for part in parts:
        text = part["text"] if isinstance(part, dict) else str(part)
        run = paragraph.add_run(text)
        if isinstance(part, dict):
            set_run_font(
                run,
                size=part.get("size", default_size),
                bold=part.get("bold"),
                color=part.get("color"),
                font=part.get("font", "Calibri"),
            )
        else:
            set_run_font(run, size=default_size)


def add_para(doc, parts, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if isinstance(parts, str):
        parts = [{"text": parts}]
    write_runs(p, parts)
    return p


def add_bullet(doc, parts):
    return add_para(doc, parts, style="List Bullet")


def add_step(doc, title, detail=None):
    p = doc.add_paragraph(style="List Number")
    write_runs(p, [{"text": title, "bold": True, "color": BLUE}])
    if detail:
        write_runs(p, [{"text": f"：{detail}"}])
    return p


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                cell = row.cells[idx]
                cell.width = Inches(width)
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table_text(table, header_rows=0, label_cols=0):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, bold=(i < header_rows or j < label_cols))
            set_cell_margins(cell)
            if i < header_rows or j < label_cols:
                set_cell_shading(cell, LIGHT_BLUE)
    add_border(table)


def add_key_table(doc):
    table = doc.add_table(rows=4, cols=2)
    set_table_width(table, [1.65, 4.65])
    rows = [
        ("文档对象", "普通用户、应用配置人员、内部使用人员"),
        ("中转站地址", "https://ai.muling.store"),
        ("接口地址", "https://ai.muling.store/v1"),
        ("默认模型", "gpt-5.5"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        p = table.cell(i, 1).paragraphs[0]
        p.clear()
        write_runs(p, [{"text": v, "bold": v in {"https://ai.muling.store", "https://ai.muling.store/v1", "gpt-5.5"}, "color": RED if v != rows[0][1] else None}])
    style_table_text(table, label_cols=1)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    add_border(table, color="D9E2EF", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=9.5, color="1F3A5F", font="Consolas")
    return table


def add_warning_box(doc, parts):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    add_border(table, color="E6B8B7", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FDE9D9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.clear()
    write_runs(p, parts)
    return table


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    cover = doc.add_table(rows=1, cols=1)
    set_table_width(cover, [6.3])
    add_border(cover, color="D9E2EF", size="8")
    set_cell_shading(cover.cell(0, 0), "F4F8FC")
    set_cell_margins(cover.cell(0, 0), top=260, bottom=260, start=260, end=260)
    p = cover.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_runs(p, [{"text": "NewAPI 第三方 API 接口配置 OpenClaw 小龙虾用户操作手册", "size": 21, "bold": True, "color": TITLE_BLUE}])
    p2 = cover.cell(0, 0).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_runs(p2, [
        {"text": "适用于 ", "size": 11, "color": MUTED},
        {"text": "https://ai.muling.store", "size": 11, "bold": True, "color": RED},
        {"text": "；默认模型：", "size": 11, "color": MUTED},
        {"text": "gpt-5.5", "size": 11, "bold": True, "color": RED},
    ])

    add_para(doc, "")
    add_key_table(doc)

    doc.add_heading("一、使用前准备", level=1)
    add_para(doc, "本文用于指导用户在 OpenClaw 小龙虾中使用 NewAPI 中转站提供的第三方 API 接口。你不需要自己配置上游渠道，只需要拿到中转站账号和 API Key 后，在客户端中填写接口地址、密钥和模型名称。")
    add_para(doc, [
        {"text": "说明：", "bold": True, "color": RED},
        {"text": "如果你的电脑尚未安装 WSL 或小龙虾 OpenClaw，请先浏览本文最后的“附录：Windows 中使用 WSL 安装小龙虾 OpenClaw”。"},
    ])
    add_bullet(doc, [{"text": "中转站地址："}, {"text": "https://ai.muling.store", "bold": True, "color": RED}])
    add_bullet(doc, [{"text": "OpenAI 兼容接口地址："}, {"text": "https://ai.muling.store/v1", "bold": True, "color": RED}])
    add_bullet(doc, [{"text": "默认模型统一选择："}, {"text": "gpt-5.5", "bold": True, "color": RED}])
    add_bullet(doc, "API Key 由中转站后台生成，用于调用第三方 API 接口，请勿在聊天、截图或公共文档中泄露。")

    add_warning_box(doc, [
        {"text": "重要：", "bold": True, "color": RED},
        {"text": "本文档中的 API Key 示例均为占位符。实际使用时请替换为你自己的密钥，并妥善保管。"},
    ])

    doc.add_heading("二、获取 NewAPI API Key", level=1)
    add_step(doc, "登录中转站", "打开 https://ai.muling.store，输入你的账号密码完成登录。")
    add_step(doc, "进入令牌页面", "在个人中心、令牌管理或 API Key 页面中新建令牌。")
    add_step(doc, "复制 API Key", "复制以 sk- 开头的密钥，并保存到本地安全位置。")
    add_step(doc, "确认可用模型", "在模型列表中确认 gpt-5.5 可用；如果看不到该模型，请联系管理员开通。")

    doc.add_heading("三、在 OpenClaw 中配置 NewAPI 第三方 API 接口", level=1)
    add_para(doc, "OpenClaw 需要填写提供方、Base URL、API Key 和默认模型。这里的 Base URL 和 API Key 都来自 NewAPI 中转站，不是 OpenAI 官方接口。")
    cfg = doc.add_table(rows=5, cols=2)
    set_table_width(cfg, [2.0, 4.3])
    cfg_rows = [
        ("提供方类型", "OpenAI Compatible / OpenAI 兼容"),
        ("Base URL", "https://ai.muling.store/v1"),
        ("API Key", "填写你在中转站生成的 sk- 密钥"),
        ("默认模型", "gpt-5.5"),
        ("模型写法", "如 OpenClaw 要求 provider/model，可填写 openai/gpt-5.5"),
    ]
    for i, (k, v) in enumerate(cfg_rows):
        cfg.cell(i, 0).text = k
        p = cfg.cell(i, 1).paragraphs[0]
        p.clear()
        is_key = v in {"https://ai.muling.store/v1", "gpt-5.5"} or "openai/gpt-5.5" in v
        write_runs(p, [{"text": v, "bold": is_key, "color": RED if is_key else None}])
    style_table_text(cfg, label_cols=1)

    doc.add_heading("四、推荐命令行配置方式", level=1)
    add_para(doc, "如果你的 OpenClaw 环境支持命令行配置，可参考以下流程。不同版本命令名称可能略有差异，以实际提示为准。")
    add_code_block(doc, [
        "openclaw onboard --install-daemon",
        "# Provider 选择：OpenAI Compatible",
        "# Base URL 填写：https://ai.muling.store/v1",
        "# API Key 填写：sk-xxxxxxxx",
        "openclaw models set openai/gpt-5.5",
        "openclaw models status",
    ])
    add_warning_box(doc, [
        {"text": "关键配置：", "bold": True, "color": RED},
        {"text": "Base URL 必须填写 "},
        {"text": "https://ai.muling.store/v1", "bold": True, "color": RED},
        {"text": "，默认模型必须选择 "},
        {"text": "gpt-5.5", "bold": True, "color": RED},
        {"text": "。"},
    ])

    doc.add_heading("五、接口连通性测试", level=1)
    add_para(doc, "配置完成后，建议先做一次接口测试。将 sk-xxxxxxxx 替换为你的真实 API Key。")
    add_code_block(doc, [
        "curl https://ai.muling.store/v1/chat/completions \\",
        "  -H \"Authorization: Bearer sk-xxxxxxxx\" \\",
        "  -H \"Content-Type: application/json\" \\",
        "  -d '{",
        "    \"model\": \"gpt-5.5\",",
        "    \"messages\": [{\"role\": \"user\", \"content\": \"请回复：连接成功\"}],",
        "    \"stream\": false",
        "  }'",
    ])
    add_para(doc, "如果返回正常回复，说明中转站、API Key 和模型配置均已生效。")

    doc.add_heading("六、常见问题处理", level=1)
    faq = doc.add_table(rows=5, cols=3)
    set_table_width(faq, [1.45, 2.2, 2.65])
    for j, h in enumerate(["现象", "可能原因", "处理方法"]):
        faq.cell(0, j).text = h
    rows = [
        ("401 Unauthorized", "API Key 错误、过期或未启用", "重新生成 API Key，并确认复制完整。"),
        ("model not found", "模型名错误或账号未开通", "确认模型名为 gpt-5.5，或联系管理员开通。"),
        ("404 Not Found", "Base URL 写错", "确认填写 https://ai.muling.store/v1。"),
        ("无回复或超时", "网络或上游渠道异常", "稍后重试；仍失败时联系管理员检查渠道。"),
    ]
    for i, row in enumerate(rows, 1):
        for j, value in enumerate(row):
            faq.cell(i, j).text = value
    style_table_text(faq, header_rows=1)

    doc.add_heading("七、用户检查清单", level=1)
    for item in [
        "已能登录 https://ai.muling.store。",
        "已生成并保存 API Key。",
        "如需本机安装，小龙虾 OpenClaw 已安装完成，openclaw 命令可正常执行。",
        "Base URL 已填写为 https://ai.muling.store/v1。",
        "默认模型已选择 gpt-5.5。",
        "已完成一次简单对话测试。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("八、安全提醒", level=1)
    add_bullet(doc, [{"text": "API Key 等同于账号调用凭证，", "bold": True, "color": RED}, {"text": "请不要发到群聊、截图或公开文档中。"}])
    add_bullet(doc, "如怀疑 API Key 泄露，请立即在中转站后台删除或重置。")
    add_bullet(doc, "多人共用设备时，建议不要勾选记住密钥或自动保存配置。")

    doc.add_heading("九、推荐配置摘要", level=1)
    add_code_block(doc, [
        "NewAPI URL: https://ai.muling.store",
        "Base URL:   https://ai.muling.store/v1",
        "Model:      gpt-5.5",
        "OpenClaw:   openai/gpt-5.5",
    ])

    doc.add_page_break()
    doc.add_heading("附录：Windows 中使用 WSL 安装小龙虾 OpenClaw", level=1)
    add_para(doc, "本附录仅适用于电脑上尚未安装 OpenClaw 小龙虾的用户。如果你已经能打开 OpenClaw 或已经能执行 openclaw 命令，可以跳过本附录，直接按前文配置第三方 API 接口。")
    add_warning_box(doc, [
        {"text": "关键要求：", "bold": True, "color": RED},
        {"text": "请使用 Windows Terminal 或 PowerShell 以管理员身份执行 WSL 安装命令；安装完成后通常需要重启电脑。"},
    ])

    doc.add_heading("A.1 安装 WSL", level=2)
    add_step(doc, "打开 PowerShell 管理员窗口", "在开始菜单搜索 PowerShell，右键选择“以管理员身份运行”。")
    add_step(doc, "执行安装命令", "安装默认 Ubuntu 发行版。")
    add_code_block(doc, ["wsl --install"])
    add_step(doc, "重启电脑", "安装完成后按提示重启 Windows。")
    add_step(doc, "创建 Linux 用户", "首次打开 Ubuntu 时，按提示创建用户名和密码。")

    doc.add_heading("A.2 更新 Ubuntu 环境", level=2)
    add_para(doc, "进入 Ubuntu 终端后，执行以下命令更新软件源：")
    add_code_block(doc, ["sudo apt update", "sudo apt upgrade -y"])

    doc.add_heading("A.3 安装小龙虾 OpenClaw", level=2)
    add_para(doc, "请优先使用 OpenClaw 官方提供的最新安装命令。如果团队已提供固定安装脚本，请以团队脚本为准。")
    add_code_block(doc, [
        "# 示例：按官方文档或团队文档提供的安装命令执行",
        "# 安装完成后执行初始化",
        "openclaw onboard --install-daemon",
    ])
    add_warning_box(doc, [
        {"text": "关键：", "bold": True, "color": RED},
        {"text": "如果系统提示找不到 openclaw 命令，请关闭并重新打开 Ubuntu 终端，或确认安装目录已加入 PATH。"},
    ])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    write_runs(footer, [{"text": "NewAPI / OpenClaw 小龙虾用户操作手册", "size": 9, "color": MUTED}])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
